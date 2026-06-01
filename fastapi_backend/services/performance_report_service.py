"""
性能测试报告生成器

参照京通真实报告格式，生成包含完整指标和图表的 Word 文档
"""

import io
import os
import tempfile
import httpx
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from fastapi_backend.services.performance_charts import generate_all_charts


def __safe_err_rate(val):
    """安全格式化错误率"""
    try:
        v = float(val or 0)
        if v != v or v > 1e9 or v < 0:
            return "N/A"
        return f"{v:.2f}"
    except (ValueError, TypeError):
        return "N/A"


def _is_placeholder_api_key(key):
    if not key:
        return True
    key_lower = key.lower()
    placeholders = [
        "your_model_api_key_here",
        "your-openai-api-key",
        "your_api_key",
        "your-api-key",
        "sk-your",
        "change_me",
        "changeme",
        "placeholder",
    ]
    return any(p in key_lower for p in placeholders)


async def _call_ai_analysis(scenarios: List[Dict], summary: Dict, test_env: Dict) -> str:
    """调用 AI 生成优化建议"""
    from sqlalchemy import select
    from fastapi_backend.models.models import AIConfig
    from fastapi_backend.core.database import AsyncSessionLocal

    try:
        async with AsyncSessionLocal() as db:
            result = await db.execute(select(AIConfig).where(AIConfig.is_active))
            config = result.scalar_one_or_none()
    except Exception:
        config = None

    api_key = None
    base_url = None
    model = None

    if config:
        from fastapi_backend.utils.encryption import decrypt, DecryptionError
        try:
            api_key = decrypt(config.api_key)
        except (DecryptionError, Exception):
            api_key = config.api_key
        base_url = config.base_url
        model = config.model
    else:
        from fastapi_backend.core.config import settings

        if settings.AI_API_KEY and not _is_placeholder_api_key(settings.AI_API_KEY):
            api_key = settings.AI_API_KEY
            base_url = settings.AI_BASE_URL
            model = settings.AI_MODEL

    if not api_key or not model:
        return _build_offline_suggestions(scenarios, summary)

    scenario_text = ""
    for s in scenarios:
        err_rate = s.get("error_rate", 0) or 0
        scenario_text += (
            f"- {s.get('name', '')}: TPS={s.get('actual_qps', 'N/A')}, "
            f"平均={s.get('avg_ms', 'N/A')}ms, P50={s.get('p50_ms', 'N/A')}ms, "
            f"P90={s.get('p90_ms', 'N/A')}ms, P95={s.get('p95_ms', 'N/A')}ms, "
            f"P99={s.get('p99_ms', 'N/A')}ms, 错误率={err_rate:.2f}%, "
            f"请求数={s.get('total_requests', 0) or 0}, 失败={s.get('failed_requests', 0) or 0}\n"
        )

    prompt = f"""你是资深性能测试专家。请针对以下压测结果给出专业的优化建议报告（300-500字，中文，分点建议）。

测试环境: {test_env.get("env_name", "生产环境")}
测试域名: {test_env.get("domain", "")}

各场景性能数据:
{scenario_text}

整体统计:
- 总请求: {summary.get("total_requests", "N/A")}
- 总失败: {summary.get("total_failed", "N/A")}
- 整体错误率: {summary.get("overall_error_rate", "N/A")}%

请按以下结构输出:
1. 整体性能评估（优秀/良好/一般/较差）
2. 瓶颈分析（P95/P99差距、标准差分析）
3. 失败原因分析
4. 优化建议（至少3条可操作的）

用中文，简洁专业，适合直接放入正式报告。"""

    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                f"{base_url}/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": model,
                    "messages": [
                        {
                            "role": "system",
                            "content": "你是资深性能测试专家，输出专业、简洁、可直接放入正式报告的分析。",
                        },
                        {"role": "user", "content": prompt},
                    ],
                    "max_tokens": 1000,
                    "temperature": 0.3,
                },
            )
            if resp.status_code == 200:
                data = resp.json()
                return data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
    except Exception:
        pass

    return _build_offline_suggestions(scenarios, summary)


def _build_offline_suggestions(scenarios: List[Dict], summary: Dict) -> str:
    """离线模式生成建议"""
    lines = ["（以下为基于规则自动生成的建议）", ""]
    passed = sum(1 for s in scenarios if s.get("result", "") in ("通过", "pass", "Pass", "PASS", "PASSED", "成功"))
    failed = len(scenarios) - passed

    if failed == 0:
        lines.append("1. 整体性能评估：优秀 - 所有场景均通过测试。")
    elif failed <= len(scenarios) * 0.3:
        lines.append(f"1. 整体性能评估：良好 - {passed}/{len(scenarios)} 场景通过，部分场景需优化。")
    else:
        lines.append(f"1. 整体性能评估：需改进 - 仅 {passed}/{len(scenarios)} 场景通过，建议全面排查。")

    idx = 2
    for s in scenarios:
        name = s.get("name", "")
        p95 = s.get("p95_ms", 0) or 0
        p99 = s.get("p99_ms", 0) or 0
        err_rate = s.get("error_rate", 0) or 0
        if p99 > 0 and p95 > 0 and (p99 - p95) > p95 * 0.5:
            lines.append(f"{idx}. 瓶颈分析：{name} 的 P99({p99}ms) 远高于 P95({p95}ms)，存在长尾延迟，建议排查慢请求。")
            idx += 1
        if err_rate > 1:
            lines.append(f"{idx}. 失败分析：{name} 错误率 {err_rate:.2f}%，超过 1% 标准，建议检查服务端日志。")
            idx += 1

    lines.append(f"{idx}. 优化建议：")
    lines.append("   - 检查数据库慢查询并添加索引")
    lines.append("   - 对高耗时接口考虑增加缓存层")
    lines.append("   - 检查连接池配置是否合理")
    lines.append("   - 对频繁超时的接口考虑异步处理或熔断")

    return "\n".join(lines)


def generate_performance_report(
    report_name: str,
    test_env: Dict[str, Any],
    scenarios: List[Dict[str, Any]],
    summary: Dict[str, Any] = None,
    author: str = "",
    error_types: Dict[str, int] = None,
    ai_suggestions: str = "",
    env_config: Dict[str, Any] = None,
    rt_distribution: Dict[str, int] = None,
    throughput_trend: List[Dict] = None,
    status_distribution: Dict[str, int] = None,
) -> bytes:
    """
    生成完整性能测试报告 (docx)

    Args:
        report_name: 报告名称
        test_env: {"domain": "https://api.example.com", "env_name": "生产环境"}
        scenarios: [{name, url, method, target_qps, actual_qps, result,
                     avg_ms, p50_ms, p90_ms, p95_ms, p99_ms, stddev_ms, max_ms, min_ms,
                     concurrency, threads, ramp_up, loops, duration,
                     error_rate, total_requests, failed_requests,
                     test_start, test_end, analysis}, ...]
        summary: {total_requests, total_failed, overall_error_rate, overall_result, notes}
        author: 报告作者
        error_types: {"连接超时": 5, "HTTP 500": 3, ...}
        ai_suggestions: AI生成的优化建议文本
        env_config: {cpu, memory, os, jdk, jmeter_version, network, ...}
        rt_distribution: {"<10ms": 100, "10-50ms": 500, ...}
        throughput_trend: [{"t": 0, "tps": 50.2}, ...]
        status_distribution: {"200": 950, "500": 50}

    Returns: docx bytes
    """
    doc = Document()

    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _set_cell(cell, text, bold=False, color=None, align=None):
        cell.text = ""
        p = cell.paragraphs[0]
        if align:
            p.alignment = align
        # 过滤 NaN/Infinity
        if isinstance(text, float):
            if not (text > -1e9 and text < 1e9) or text != text:
                text = "N/A"
        run = p.add_run(str(text))
        run.bold = bold
        if color:
            run.font.color.rgb = color

    # ========== 封面 ==========
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(report_name)
    title_run.font.size = Pt(22)
    title_run.bold = True

    doc.add_paragraph()

    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("文件编号", f"TM-PERF-{datetime.now().strftime('%Y%m%d')}"),
        ("文件状态", "正式发布"),
        ("当前版本", "V1.0"),
        ("拟制人", author or "TestMaster"),
        ("日期", datetime.now().strftime("%Y/%m/%d")),
        ("测试环境", test_env.get("env_name", "生产环境")),
    ]
    for i, (key, val) in enumerate(meta_data):
        _set_cell(meta_table.cell(i, 0), key, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(meta_table.cell(i, 1), val, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ========== 1. 背景 ==========
    doc.add_heading("1. 背景", level=1)
    doc.add_paragraph(f"测试域名：{test_env.get('domain', '')}")
    doc.add_paragraph(f"本报告对 {len(scenarios)} 个场景/接口进行了性能压测，评估系统的吞吐能力、响应时间和稳定性。")

    # ========== 2. 测试概述 ==========
    doc.add_heading("2. 测试概述", level=1)

    doc.add_heading("2.1 测试目的", level=2)
    doc.add_paragraph("确定系统在生产环境下的并发承载能力、响应时间表现，检验系统瓶颈，获取各接口的处理能力数据。")

    doc.add_heading("2.2 压测目标环境", level=2)
    doc.add_paragraph(f"环境名称：{test_env.get('env_name', '生产环境')}")
    doc.add_paragraph(f"目标域名：{test_env.get('domain', '')}")

    # ========== 2.3 测试环境配置 ==========
    if env_config:
        doc.add_heading("2.3 测试环境配置", level=2)
        env_table = doc.add_table(rows=len(env_config) + 1, cols=2)
        env_table.style = "Table Grid"
        env_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(["配置项", "参数值"]):
            _set_cell(env_table.cell(0, j), h, bold=True)
        for i, (k, v) in enumerate(env_config.items()):
            _set_cell(env_table.cell(i + 1, 0), k, bold=True)
            _set_cell(env_table.cell(i + 1, 1), str(v))
        sec_offset = 1
    else:
        sec_offset = 0

    sec_num = 3 + sec_offset

    doc.add_heading(f"2.{sec_num} 压测场景所含接口", level=2)
    api_table = doc.add_table(rows=len(scenarios) + 1, cols=5)
    api_table.style = "Table Grid"
    api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["序号", "接口 URL", "方法", "目标 QPS", "场景名称"]):
        _set_cell(api_table.cell(0, j), h, bold=True)
    for i, s in enumerate(scenarios):
        _set_cell(api_table.cell(i + 1, 0), str(i + 1))
        _set_cell(api_table.cell(i + 1, 1), s.get("url", ""))
        _set_cell(api_table.cell(i + 1, 2), s.get("method", "GET"))
        _set_cell(api_table.cell(i + 1, 3), str(s.get("target_qps", "未设定")))
        _set_cell(api_table.cell(i + 1, 4), s.get("name", ""))

    doc.add_heading(f"2.{sec_num + 1} 测试指标/参考依据", level=2)
    indicator_table = doc.add_table(rows=7, cols=3)
    indicator_table.style = "Table Grid"
    for j, h in enumerate(["序号", "指标项", "标准"]):
        _set_cell(indicator_table.cell(0, j), h, bold=True)
    indicators = [
        ("1", "错误率", "1% 以内"),
        ("2", "平均响应时间", "核心接口 < 500ms"),
        ("3", "P95 响应时间", "5s 以内"),
        ("4", "P99 响应时间", "10s 以内"),
        ("5", "吞吐量 (QPS)", "核心接口峰值吞吐量达到项目要求"),
        ("6", "并发用户数", "支持目标并发数且无系统性崩溃"),
    ]
    for i, (no, name, std) in enumerate(indicators):
        _set_cell(indicator_table.cell(i + 1, 0), no)
        _set_cell(indicator_table.cell(i + 1, 1), name)
        _set_cell(indicator_table.cell(i + 1, 2), std)

    doc.add_heading(f"2.{sec_num + 2} 术语/缩略语", level=2)
    term_table = doc.add_table(rows=8, cols=3)
    term_table.style = "Table Grid"
    for j, h in enumerate(["术语", "全称", "含义"]):
        _set_cell(term_table.cell(0, j), h, bold=True)
    terms = [
        ("TPS", "Transaction Per Second", "每秒处理事务数"),
        ("QPS", "Queries Per Second", "每秒处理请求数"),
        ("并发数", "Concurrency", "同一时刻与服务器交互的在线用户数"),
        (
            "P50/P90/P95/P99",
            "Percentile",
            "将响应时间排序后对应百分位上的值，反映延迟分布",
        ),
        ("标准差", "Standard Deviation", "响应时间的离散程度，越小越稳定"),
        ("错误率", "Error Rate", "返回错误状态码的请求数 / 总请求数"),
        ("VUM", "Virtual User Minute", "1个虚拟用户执行1分钟的压测消耗1个VUM"),
    ]
    for i, (t, f_n, m) in enumerate(terms):
        _set_cell(term_table.cell(i + 1, 0), t)
        _set_cell(term_table.cell(i + 1, 1), f_n)
        _set_cell(term_table.cell(i + 1, 2), m)

    doc.add_page_break()

    # ========== 3. 性能测试过程 ==========
    doc.add_heading("3. 性能测试过程", level=1)
    doc.add_paragraph(f"{test_env.get('env_name', '生产环境')}进行压测")

    for idx, s in enumerate(scenarios):
        status = s.get("result", "通过")
        status_icon = "通过" if status in ("通过", "pass") else "失败"

        doc.add_heading(f"3.{idx + 1} 【{status_icon}】场景：{s.get('name', '')}", level=2)

        t = doc.add_table(rows=15, cols=2)
        t.style = "Table Grid"

        result_data = [
            ("测试结果", status_icon),
            ("吞吐量 (QPS)", f"{s.get('actual_qps', 'N/A')}"),
            ("报错率", f"{s.get('error_rate', 0):.2f}%"),
            ("平均响应时间", f"{s.get('avg_ms', 'N/A')}ms"),
            ("P50 响应时间", f"{s.get('p50_ms', 'N/A')}ms"),
            ("P90 响应时间", f"{s.get('p90_ms', 'N/A')}ms"),
            ("P95 响应时间", f"{s.get('p95_ms', 'N/A')}ms"),
            ("P99 响应时间", f"{s.get('p99_ms', 'N/A')}ms"),
            ("标准差", f"{s.get('stddev_ms', 'N/A')}ms"),
            ("最大响应时间", f"{s.get('max_ms', 'N/A')}ms"),
            ("最小响应时间", f"{s.get('min_ms', 'N/A')}ms"),
            ("并发量", str(s.get("concurrency", "N/A"))),
            ("线程数", str(s.get("threads", "N/A"))),
            ("循环次数", str(s.get("loops", "N/A"))),
            ("持续时间", f"{s.get('duration', 'N/A')}秒"),
        ]

        for i, (k, v) in enumerate(result_data):
            _set_cell(t.cell(i, 0), k, bold=True)
            if k == "测试结果":
                if status_icon == "通过":
                    _set_cell(t.cell(i, 1), v, bold=True, color=RGBColor(0x10, 0xB9, 0x81))
                else:
                    _set_cell(t.cell(i, 1), v, bold=True, color=RGBColor(0xEF, 0x44, 0x44))
            else:
                _set_cell(t.cell(i, 1), v)

        doc.add_paragraph()
        doc.add_paragraph(f"测试时间：{s.get('test_start', '')} 到 {s.get('test_end', '')}")

        analysis = s.get("analysis", "")
        if analysis:
            doc.add_paragraph(f"数据分析：{analysis}")

    # ========== 3.X 错误详情分析 ==========
    if error_types and sum(error_types.values()) > 0:
        doc.add_heading(f"3.{len(scenarios) + 1} 错误详情分析", level=2)
        doc.add_paragraph(f"本次压测共出现 {sum(error_types.values())} 个错误，详细分布如下：")

        err_table = doc.add_table(rows=len(error_types) + 1, cols=3)
        err_table.style = "Table Grid"
        err_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(["错误类型", "出现次数", "占比"]):
            _set_cell(err_table.cell(0, j), h, bold=True)

        total_err = sum(error_types.values())
        sorted_errs = sorted(error_types.items(), key=lambda x: x[1], reverse=True)
        for i, (etype, cnt) in enumerate(sorted_errs):
            _set_cell(err_table.cell(i + 1, 0), etype[:80])
            _set_cell(err_table.cell(i + 1, 1), str(cnt))
            _set_cell(err_table.cell(i + 1, 2), f"{cnt / total_err * 100:.1f}%")

    doc.add_page_break()

    # ========== 4. 性能图表 ==========
    doc.add_heading("4. 性能图表", level=1)

    try:
        charts = generate_all_charts(
            scenarios,
            rt_distribution=rt_distribution,
            throughput_trend=throughput_trend,
            status_distribution=status_distribution,
            error_types=error_types,
        )
        chart_titles = {
            "tps_compare": "图1: 各接口TPS吞吐量对比",
            "response_time_multimetric": "图2: 各接口响应时间多指标对比 (Avg/P50/P90/P95/P99)",
            "rt_distribution": "图3: 响应时间区间分布",
            "throughput_trend": "图4: TPS吞吐量趋势",
            "status_distribution": "图5: HTTP状态码分布",
            "error_type_pie": "图6: 错误类型分布",
        }

        tmpdir = tempfile.mkdtemp()
        try:
            for chart_key, title in chart_titles.items():
                if chart_key in charts:
                    doc.add_heading(title, level=2)
                    img_path = os.path.join(tmpdir, f"{chart_key}.png")
                    with open(img_path, "wb") as f:
                        f.write(charts[chart_key].getvalue())
                    section = doc.sections[0]
                    page_width = section.page_width - section.left_margin - section.right_margin
                    img_width = min(Inches(5.5), page_width)
                    doc.add_picture(img_path, width=img_width)
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
        finally:
            import shutil

            shutil.rmtree(tmpdir, ignore_errors=True)
    except Exception as e:
        doc.add_paragraph(f"(图表生成失败: {e})。请检查服务器 matplotlib 环境。")

    doc.add_page_break()

    # ========== 5. 结论 ==========
    doc.add_heading("5. 结论", level=1)
    doc.add_paragraph("接口压测结果数据如下：")

    passed = sum(1 for s in scenarios if s.get("result", "") in ("通过", "pass", "Pass", "PASS", "PASSED", "成功"))
    overall_qps = sum(float(s.get("actual_qps", 0) or 0) for s in scenarios)

    # 汇总表
    summary_table = doc.add_table(rows=len(scenarios) + 1, cols=9)
    summary_table.style = "Table Grid"
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 设置列宽避免A4溢出
    col_widths = [
        Cm(3.5),
        Cm(1.5),
        Cm(2.0),
        Cm(2.0),
        Cm(2.0),
        Cm(2.0),
        Cm(2.0),
        Cm(1.8),
        Cm(1.8),
    ]
    for j, w in enumerate(col_widths):
        summary_table.columns[j].width = w
    s_headers = [
        "场景名称",
        "并发数",
        "TPS",
        "平均(ms)",
        "P95(ms)",
        "P99(ms)",
        "标准差(ms)",
        "错误率",
        "结果",
    ]
    for j, h in enumerate(s_headers):
        _set_cell(summary_table.cell(0, j), h, bold=True)

    for i, s in enumerate(scenarios):
        vals = [
            s.get("name", ""),
            str(s.get("concurrency", "N/A")),
            str(s.get("actual_qps", "N/A")),
            str(s.get("avg_ms", "N/A")),
            str(s.get("p95_ms", "N/A")),
            str(s.get("p99_ms", "N/A")),
            str(s.get("stddev_ms", "N/A")),
            f"{__safe_err_rate(s.get('error_rate', 0))}%",
            s.get("result", "N/A"),
        ]
        for j, v in enumerate(vals):
            _set_cell(summary_table.cell(i + 1, j), v)

    doc.add_paragraph()
    total_scenarios = len(scenarios) or 1
    doc.add_paragraph(f"场景通过率：{passed}/{len(scenarios)} ({passed / total_scenarios * 100:.1f}%)")
    try:
        overall_qps_val = float(overall_qps) if overall_qps and overall_qps < 1e9 else 0
    except (ValueError, TypeError):
        overall_qps_val = 0
    doc.add_paragraph(f"合计吞吐量：{overall_qps_val:.2f} QPS")

    if summary:
        total_requests = summary.get("total_requests", "N/A")
        total_failed = summary.get("total_failed", "N/A")
        doc.add_paragraph(f"总请求数：{total_requests}")
        doc.add_paragraph(f"总失败数：{total_failed}")
        err_rate = summary.get("overall_error_rate", 0) or 0
        try:
            err_rate_str = f"{float(err_rate):.2f}%"
        except (ValueError, TypeError):
            err_rate_str = f"{err_rate}%"
        doc.add_paragraph(f"整体错误率：{err_rate_str}")

    # ========== 6. 优化建议 ==========
    doc.add_page_break()
    doc.add_heading("6. 优化建议", level=1)

    if ai_suggestions:
        import re

        for line in ai_suggestions.strip().split("\n"):
            line = line.strip()
            if line:
                # 清洗markdown标记
                cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
                cleaned = re.sub(r"__(.*?)__", r"\1", cleaned)
                cleaned = re.sub(r"`(.*?)`", r"\1", cleaned)
                cleaned = re.sub(r"^[*-]\s+", "", cleaned)
                if cleaned:
                    doc.add_paragraph(cleaned)
    else:
        doc.add_paragraph("（请配置 AI 分析以获取智能优化建议）")

    if summary and summary.get("notes"):
        doc.add_paragraph()
        doc.add_paragraph(f"补充说明：{summary['notes']}")

    # ========== 附录 ==========
    doc.add_heading("附录. 审批意见", level=1)
    doc.add_paragraph("项目经理审批意见：")
    doc.add_paragraph()
    doc.add_paragraph(f"签字：测试负责人  {author or ''}")
    doc.add_paragraph(f"日期：{datetime.now().strftime('%Y-%m-%d')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
